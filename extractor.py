#!/usr/bin/env python3
"""
extractor.py - Extract text from downloaded documents

This module processes PDFs and other document formats to extract text content.
It automatically detects whether PDFs are text-based or image-based and applies
the appropriate extraction method (direct text extraction or OCR).
"""

import os
import re
import logging
import argparse
from pathlib import Path
from typing import List, Tuple, Optional, Dict
import tempfile

import fitz  # PyMuPDF
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    pytesseract = None

from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from tqdm import tqdm
import base64
import io
import requests


def setup_logging(log_dir: str = "logs") -> None:
    """Set up logging configuration."""
    os.makedirs(log_dir, exist_ok=True)
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(f"{log_dir}/extractor.log"),
            logging.StreamHandler()
        ]
    )


def is_pdf_text_based(pdf_path: str, min_text_threshold: int = 50) -> bool:
    """
    Determine if a PDF contains extractable text or is image-based.
    
    Args:
        pdf_path: Path to PDF file
        min_text_threshold: Minimum characters needed to consider PDF text-based
        
    Returns:
        True if PDF contains extractable text, False if needs OCR
    """
    try:
        doc = fitz.open(pdf_path)
        total_text_length = 0
        
        # Check first few pages for text content
        pages_to_check = min(3, len(doc))
        
        for page_num in range(pages_to_check):
            page = doc[page_num]
            text = page.get_text()
            # Clean text and count meaningful characters
            clean_text = re.sub(r'\s+', ' ', text).strip()
            total_text_length += len(clean_text)
            
            # Early exit if we have enough text
            if total_text_length >= min_text_threshold:
                doc.close()
                return True
        
        doc.close()
        return total_text_length >= min_text_threshold
        
    except Exception as e:
        logging.warning(f"Error checking PDF text content for {pdf_path}: {str(e)}")
        return False


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text directly from a text-based PDF.
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Extracted text content
    """
    try:
        doc = fitz.open(pdf_path)
        text_content = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                # Clean up the text
                text = re.sub(r'\n\s*\n', '\n\n', text)  # Normalize line breaks
                text = re.sub(r'[ \t]+', ' ', text)      # Normalize spaces
                text_content.append(text.strip())
        
        doc.close()
        
        # Join all pages with page breaks
        full_text = '\n\n--- PAGE BREAK ---\n\n'.join(text_content)
        return full_text
        
    except Exception as e:
        logging.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
        return ""


def cloud_ocr_fallback(image: Image.Image) -> str:
    """
    Use cloud OCR services as fallback when local Tesseract unavailable.
    
    Args:
        image: PIL Image object
        
    Returns:
        Extracted text string
    """
    try:
        # Convert image to base64
        buffer = io.BytesIO()
        image.save(buffer, format='PNG')
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        
        # Try OCR.space free API (no key required for small files)
        try:
            url = "https://api.ocr.space/parse/image"
            payload = {
                'base64Image': f'data:image/png;base64,{image_base64}',
                'language': 'eng',
                'isOverlayRequired': False,
                'iscreatesearchablepdf': False,
                'issearchablepdfhidetextlayer': False
            }
            
            response = requests.post(url, data=payload, timeout=30)
            result = response.json()
            
            if result.get('IsErroredOnProcessing', False):
                logging.warning("OCR.space API error")
                return ""
            
            text = result.get('ParsedResults', [{}])[0].get('ParsedText', '')
            if text.strip():
                logging.info("Cloud OCR successful via OCR.space")
                return text.strip()
                
        except Exception as e:
            logging.warning(f"OCR.space API failed: {str(e)}")
        
        # Fallback: return empty string with helpful message
        logging.warning("Cloud OCR services temporarily unavailable")
        return "[OCR processing unavailable - please ensure high quality images]"
        
    except Exception as e:
        logging.error(f"Cloud OCR fallback failed: {str(e)}")
        return ""

def ocr_pdf_to_text(pdf_path: str, dpi: int = 300) -> str:
    """
    Extract text from image-based PDF using OCR (local or cloud fallback).
    
    Args:
        pdf_path: Path to PDF file
        dpi: DPI for image conversion (higher = better quality but slower)
        
    Returns:
        OCR extracted text content
    """
    try:
        # Check OCR availability
        if not TESSERACT_AVAILABLE:
            logging.info("Local Tesseract unavailable, using cloud OCR fallback")
        
        # Convert PDF pages to images
        logging.info(f"Converting PDF to images for OCR: {os.path.basename(pdf_path)}")
        
        # Use temporary directory to avoid memory issues with large PDFs
        with tempfile.TemporaryDirectory() as temp_dir:
            images = convert_from_path(
                pdf_path,
                dpi=dpi,
                output_folder=temp_dir,
                fmt='PNG',
                thread_count=2  # Use multiple threads for faster conversion
            )
            
            text_content = []
            
            # Process each page with OCR
            for i, image in enumerate(tqdm(images, desc="OCR Processing")):
                try:
                    if TESSERACT_AVAILABLE and pytesseract:
                        # Use local Tesseract
                        custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
                        page_text = pytesseract.image_to_string(image, config=custom_config)
                    else:
                        # Use cloud OCR fallback
                        page_text = cloud_ocr_fallback(image)
                    
                    if page_text and page_text.strip():
                        # Clean up OCR artifacts
                        page_text = re.sub(r'\n\s*\n\s*\n+', '\n\n', page_text)  # Multiple line breaks
                        page_text = re.sub(r'[^\w\s\-.,;:!?()\[\]{}"\'&@#$%/\\]', ' ', page_text)  # Remove artifacts
                        page_text = re.sub(r'\s+', ' ', page_text)  # Normalize spaces
                        text_content.append(page_text.strip())
                        
                except Exception as e:
                    logging.warning(f"OCR failed for page {i+1} of {pdf_path}: {str(e)}")
                    text_content.append(f"[OCR FAILED FOR PAGE {i+1}]")
            
            # Join all pages
            full_text = '\n\n--- PAGE BREAK ---\n\n'.join(text_content)
            
            # Log OCR statistics
            total_chars = len(full_text)
            ocr_method = "Local Tesseract" if TESSERACT_AVAILABLE else "Cloud OCR"
            logging.info(f"OCR completed for {pdf_path} using {ocr_method}: {total_chars} characters extracted from {len(images)} pages")
            
            return full_text
            
    except Exception as e:
        logging.error(f"OCR processing failed for {pdf_path}: {str(e)}")
        return f"[OCR EXTRACTION FAILED: {str(e)}]"


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        doc = Document(docx_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        full_text = '\n\n'.join(text_content)
        logging.info(f"Extracted {len(full_text)} characters from DOCX: {os.path.basename(docx_path)}")
        
        return full_text
        
    except Exception as e:
        logging.error(f"Error extracting text from DOCX {docx_path}: {str(e)}")
        return f"[DOCX EXTRACTION FAILED: {str(e)}]"


def get_output_filename(input_path: str, output_dir: str) -> str:
    """
    Generate output filename for extracted text.
    
    Args:
        input_path: Path to input file
        output_dir: Output directory
        
    Returns:
        Full path for output text file
    """
    filename = Path(input_path).stem  # Get filename without extension
    return os.path.join(output_dir, f"{filename}.txt")


def extract_document_text(file_path: str, output_dir: str) -> Tuple[bool, str, Dict[str, any]]:
    """
    Extract text from a single document file.
    
    Args:
        file_path: Path to document file
        output_dir: Directory to save extracted text
        
    Returns:
        Tuple of (success, output_path, metadata)
    """
    file_extension = Path(file_path).suffix.lower()
    output_path = get_output_filename(file_path, output_dir)
    filename = os.path.basename(file_path)
    
    # Initialize metadata
    metadata = {
        'input_file': filename,
        'file_type': file_extension,
        'extraction_method': None,
        'character_count': 0,
        'success': False,
        'error_message': None
    }
    
    try:
        text_content = ""
        
        if file_extension == '.pdf':
            # Determine if PDF is text-based or needs OCR
            if is_pdf_text_based(file_path):
                logging.info(f"Extracting text from text-based PDF: {filename}")
                text_content = extract_text_from_pdf(file_path)
                metadata['extraction_method'] = 'direct_pdf_text'
            else:
                logging.info(f"Performing OCR on image-based PDF: {filename}")
                text_content = ocr_pdf_to_text(file_path)
                metadata['extraction_method'] = 'ocr'
                
        elif file_extension == '.docx':
            logging.info(f"Extracting text from DOCX: {filename}")
            text_content = extract_text_from_docx(file_path)
            metadata['extraction_method'] = 'docx_parser'
            
        else:
            error_msg = f"Unsupported file type: {file_extension}"
            logging.warning(error_msg)
            metadata['error_message'] = error_msg
            return False, output_path, metadata
        
        # Save extracted text
        if text_content.strip():
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            metadata['character_count'] = len(text_content)
            metadata['success'] = True
            
            logging.info(f"Successfully extracted text to: {output_path}")
            return True, output_path, metadata
        else:
            error_msg = "No text content extracted"
            logging.warning(f"{error_msg} from {filename}")
            metadata['error_message'] = error_msg
            return False, output_path, metadata
            
    except Exception as e:
        error_msg = f"Extraction failed: {str(e)}"
        logging.error(f"{error_msg} for {filename}")
        metadata['error_message'] = error_msg
        return False, output_path, metadata


def process_documents(input_dir: str = "downloads", output_dir: str = "text", 
                     failed_log: str = "logs/extraction_failures.txt") -> Dict[str, any]:
    """
    Process all documents in input directory and extract text.
    
    Args:
        input_dir: Directory containing documents to process
        output_dir: Directory to save extracted text files
        failed_log: Path to log extraction failures
        
    Returns:
        Dictionary with processing statistics
    """
    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)
    
    # Find all supported document files
    supported_extensions = {'.pdf', '.docx'}
    document_files = []
    
    for file_path in Path(input_dir).glob('*'):
        if file_path.suffix.lower() in supported_extensions:
            document_files.append(str(file_path))
    
    if not document_files:
        logging.warning(f"No supported documents found in {input_dir}")
        return {'total_files': 0, 'successful': 0, 'failed': 0, 'results': []}
    
    logging.info(f"Found {len(document_files)} documents to process")
    
    # Process all documents
    results = []
    successful_count = 0
    failed_count = 0
    failed_files = []
    
    for file_path in tqdm(document_files, desc="Extracting text"):
        success, output_path, metadata = extract_document_text(file_path, output_dir)
        
        results.append(metadata)
        
        if success:
            successful_count += 1
        else:
            failed_count += 1
            failed_files.append(f"{metadata['input_file']}: {metadata.get('error_message', 'Unknown error')}")
    
    # Log failed extractions
    if failed_files:
        os.makedirs(os.path.dirname(failed_log), exist_ok=True)
        with open(failed_log, 'w', encoding='utf-8') as f:
            f.write("Failed Text Extractions:\n")
            f.write("=" * 50 + "\n")
            for failed_file in failed_files:
                f.write(f"{failed_file}\n")
        
        logging.warning(f"Logged {len(failed_files)} failed extractions to {failed_log}")
    
    # Summary statistics
    stats = {
        'total_files': len(document_files),
        'successful': successful_count,
        'failed': failed_count,
        'results': results,
        'total_characters': sum(r['character_count'] for r in results if r['success'])
    }
    
    logging.info(f"Text extraction complete: {successful_count} successful, {failed_count} failed")
    logging.info(f"Total characters extracted: {stats['total_characters']:,}")
    
    return stats


def main():
    """Main entry point for CLI usage."""
    parser = argparse.ArgumentParser(description="Extract text from documents")
    parser.add_argument("--input-dir", default="downloads",
                       help="Directory containing documents to process (default: downloads)")
    parser.add_argument("--output-dir", default="text",
                       help="Directory to save extracted text (default: text)")
    parser.add_argument("--failed-log", default="logs/extraction_failures.txt",
                       help="Path to log extraction failures (default: logs/extraction_failures.txt)")
    parser.add_argument("--log-dir", default="logs",
                       help="Directory for log files (default: logs)")
    parser.add_argument("--single-file", 
                       help="Process a single file instead of entire directory")
    
    args = parser.parse_args()
    
    # Setup logging
    setup_logging(args.log_dir)
    
    try:
        if args.single_file:
            # Process single file
            if not os.path.exists(args.single_file):
                raise FileNotFoundError(f"File not found: {args.single_file}")
            
            os.makedirs(args.output_dir, exist_ok=True)
            
            logging.info(f"Processing single file: {args.single_file}")
            success, output_path, metadata = extract_document_text(args.single_file, args.output_dir)
            
            print(f"\nExtraction Results:")
            print(f"File: {metadata['input_file']}")
            print(f"Method: {metadata['extraction_method']}")
            print(f"Success: {metadata['success']}")
            print(f"Characters: {metadata['character_count']:,}")
            
            if metadata['error_message']:
                print(f"Error: {metadata['error_message']}")
        else:
            # Process entire directory
            if not os.path.exists(args.input_dir):
                raise FileNotFoundError(f"Input directory not found: {args.input_dir}")
            
            logging.info(f"Starting text extraction from {args.input_dir}")
            stats = process_documents(args.input_dir, args.output_dir, args.failed_log)
            
            print(f"\nExtraction Summary:")
            print(f"Total files: {stats['total_files']}")
            print(f"Successful: {stats['successful']}")
            print(f"Failed: {stats['failed']}")
            print(f"Total characters: {stats['total_characters']:,}")
            
            if stats['failed'] > 0:
                print(f"Failed extractions logged to: {args.failed_log}")
                
    except Exception as e:
        logging.error(f"Error in main process: {str(e)}")
        raise


if __name__ == "__main__":
    main()